"""
Canopy Proposal Builder — v13 (mode dispatch: single / double canopy / imaging-only)

Public entry: build_proposal(data, output_path) -> Path

v13 changes from v12:
 - Mode dispatch via data["mode"] ∈ {"single_canopy", "double_canopy", "imaging_only"}
 - Single canopy mode is fully backwards-compatible with v12 (mode key is optional;
   absence of "mode" or absence of "canopies" array → single canopy / imaging fields).
 - Double canopy: shared header/customer/sales rep/quote#/terms; per-canopy
   project specs side-by-side, per-canopy CANOPY + CANOPY BRANDING sections;
   shared PRICE SIGN; shared EXISTING CANOPY REMOVAL; combined INCLUDED /
   NOT INCLUDED with sub-headers; one combined GRAND TOTAL.
 - GAS CANOPY / DIESEL CANOPY band labels in double mode.
 - Imaging-only mode: skip canopy/dimensions/lights; IMAGING + (optional)
   PRICE SIGN sections; user-input MID labor amount (no MID_Labor lookup).
 - Imaging disclaimer added under CANOPY BRANDING band when WE supply imaging
   AND under IMAGING band in imaging-only mode (same condition).
 - Filename suffix change is handled by the caller (canopy_configurator.py).
"""

from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

APEC_PALETTE = {
    "band_bg": RGBColor(0x2B, 0x33, 0x40),
    "grand_total_bg": RGBColor(0x2B, 0x33, 0x40),
    "accent": RGBColor(0xE0, 0x20, 0x20),
    "text_emphasis": RGBColor(0x2B, 0x33, 0x40),
    "customer_bg": RGBColor(0xF5, 0xF5, 0xF7),
    "subtle_bg": RGBColor(0xF5, 0xF5, 0xF7),
    "subtle_bg_alt": RGBColor(0xEE, 0xED, 0xF0),
}

GEO_PALETTE = {
    "band_bg": RGBColor(0x0F, 0x8A, 0x60),
    "grand_total_bg": RGBColor(0x0A, 0x5C, 0x42),
    "accent": RGBColor(0x0A, 0x5C, 0x42),
    "text_emphasis": RGBColor(0x0A, 0x5C, 0x42),
    "customer_bg": RGBColor(0xEA, 0xF4, 0xF0),
    "subtle_bg": RGBColor(0xEA, 0xF4, 0xF0),
    "subtle_bg_alt": RGBColor(0xDF, 0xEE, 0xEC),
}

PALETTE = APEC_PALETTE


def _H(rgb):
    return bytes(rgb).hex().upper()


COMPANIES = {
    "APEC": {
        "name": "APEC Imaging and Canopies",
        "tagline": "Petroleum Canopy Specialists",
        "address": "4732 N Royal Atlanta Drive, Suite E, Tucker, GA 30084",
        "phone": "855-444-APEC",
        "email": "sales@TheAPECgroup.com",
        "website": "www.TheAPECgroup.com",
        "logo": "Apec Imaging Logo.jpg",
        "signature_label": "APEC Authorized Representative & Date",
        "palette": APEC_PALETTE,
    },
    "GEO": {
        "name": "GEO Canopies",
        "tagline": "Petroleum Canopy Specialists",
        "address": "40 Lyerly Street, Houston, TX 77022",
        "phone": "844-GEO-4040",
        "email": "sales@GeoPetroleum.com",
        "website": "www.GeoPetroleum.com",
        "logo": "GEO Canopies logo.jpg",
        "signature_label": "GEO Authorized Representative & Date",
        "palette": GEO_PALETTE,
    },
}

USABLE_W = 7.3
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)

# v13 imaging disclaimer (shown under CANOPY BRANDING band when WE supply,
# and under IMAGING band in imaging-only mode when WE supply).
IMAGING_DISCLAIMER = (
    "Imaging material is a direct pass through with no additional mark up. "
    "Full payment is due when the material is purchased."
)

CANOPY_BUILD_SPECS = [
    "Clearance: 16.5 ft",
    "Columns: 10″ × 10″ square",
    "Decking: white embossed, 20-gauge",
    "Perimeter gutter",
    "Scuppers",
    "Anchor bolts",
]

_TCPR_ORDER = [
    "cnfStyle",
    "tcW",
    "gridSpan",
    "hMerge",
    "vMerge",
    "tcBorders",
    "shd",
    "noWrap",
    "tcMar",
    "textDirection",
    "tcFitText",
    "vAlign",
    "hideMark",
    "headers",
    "cellIns",
    "cellDel",
    "cellMerge",
    "tcPrChange",
]


def _insert_tc_pr_ordered(tc_pr, new_element):
    new_name = new_element.tag.split("}")[-1]
    new_pos = _TCPR_ORDER.index(new_name)
    for child in list(tc_pr):
        cname = child.tag.split("}")[-1]
        if cname == new_name:
            tc_pr.remove(child)
            break
    for i, child in enumerate(tc_pr):
        cname = child.tag.split("}")[-1]
        if cname in _TCPR_ORDER and _TCPR_ORDER.index(cname) > new_pos:
            tc_pr.insert(i, new_element)
            return
    tc_pr.append(new_element)


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_tc_pr_ordered(tc_pr, shd)


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tbl_pr.append(borders)


def _set_fixed_layout(table, col_widths_inches):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for ci, w in enumerate(col_widths_inches):
        table.columns[ci].width = Inches(w)
    for row in table.rows:
        for ci, w in enumerate(col_widths_inches):
            row.cells[ci].width = Inches(w)


def _set_cell_margins(cell, top=40, left=80, bottom=40, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    _insert_tc_pr_ordered(tc_pr, mar)


def _format_run(run, *, size=10, bold=False, italic=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_para(
    cell_or_doc,
    text="",
    *,
    size=10,
    bold=False,
    italic=False,
    color=None,
    align=None,
    space_after=2,
    space_before=0,
):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        _format_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def _clear_default_para(container):
    if container.paragraphs and not container.paragraphs[0].text:
        p = container.paragraphs[0]
        p._element.getparent().remove(p._element)


def _add_cell_bottom_border(cell, sz=8, color=None):
    if color is None:
        color = _H(PALETTE["band_bg"])
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:color"), color)
    borders.append(b)
    _insert_tc_pr_ordered(tc_pr, borders)


def _money(n):
    return f"${n:,.2f}"


# v13 mode helpers ---------------------------------------------------------


def _get_mode(data):
    m = data.get("mode")
    if m in ("single_canopy", "double_canopy", "imaging_only"):
        return m
    if "canopies" in data and isinstance(data["canopies"], list):
        return "double_canopy"
    if "imaging" in data:
        return "imaging_only"
    return "single_canopy"


def _shared_block(data):
    if "shared" in data and isinstance(data["shared"], dict):
        s = dict(data["shared"])
        s.setdefault("mid", {})
        s.setdefault("demo", {})
        s.setdefault("tax_rate", 0.0825)
        s.setdefault("items_not_included", [])
        return s
    return {
        "mid": data.get(
            "mid", {"include_material": False, "include_labor": False, "brand": None}
        ),
        "demo": data.get("demo", {"include": False, "retail": 0, "cost": 0}),
        "tax_rate": data.get("pricing", {}).get("tax_rate", 0.0825),
        "items_not_included": data.get("pricing", {}).get("items_not_included", []),
    }


def _canopies_list(data):
    if "canopies" in data and isinstance(data["canopies"], list):
        return data["canopies"]
    c = dict(data.get("canopy", {}))
    p = dict(data.get("pricing", {}))
    c["pricing"] = p
    return [c]


def _mid_state(shared_or_data):
    if (
        "mid" in shared_or_data
        and isinstance(shared_or_data["mid"], dict)
        and (
            "include_material" in shared_or_data["mid"]
            or "include_labor" in shared_or_data["mid"]
            or "include" in shared_or_data["mid"]
        )
    ):
        m = shared_or_data["mid"]
    else:
        m = _shared_block(shared_or_data)["mid"]
    if "include_material" in m or "include_labor" in m:
        return (
            bool(m.get("include_material")),
            bool(m.get("include_labor")),
            m.get("brand"),
        )
    inc = bool(m.get("include"))
    return inc, inc, m.get("brand")


def _customer_supplied_imaging(canopy):
    if "customer_supplied_imaging" in canopy:
        return bool(canopy["customer_supplied_imaging"]) and bool(canopy.get("branded"))
    pricing = canopy.get("pricing", {})
    return bool(canopy.get("branded")) and pricing.get("brand_imaging", 0) == 0


def _canopy_band_label(canopy, label_prefix=None):
    prefix = label_prefix or "CANOPY"
    return f"{prefix}  —  Material, Steel and Installation"


# Per-canopy & combined totals ---------------------------------------------


def _canopy_totals(canopy):
    p = canopy.get("pricing", {})
    gp = p.get("gp_rate", 0)
    canopy_mat_marked = (
        p.get("acm", 0) + p.get("steel", 0) + p.get("decking", 0) + p.get("misc", 0)
    ) * (1 + gp)
    canopy_mat = (
        canopy_mat_marked
        + p.get("steel_secondary", 0)
        + p.get("lights", 0)
        + p.get("flood_lights", 0)
    )
    canopy_inst_marked = p.get("base_labor", 0) * (1 + gp)
    canopy_inst = (
        canopy_inst_marked
        + p.get("two_disp_labor_add", 0)
        + p.get("branded_labor_add", 0)
        + p.get("travel_adder_retail", 0)
    )
    branding_mat = p.get("brand_imaging", 0) + p.get("shipping", 0)
    branding_inst = 0
    return {
        "canopy_mat": canopy_mat,
        "canopy_inst": canopy_inst,
        "branding_mat": branding_mat,
        "branding_inst": branding_inst,
    }


def _imaging_totals(data):
    img = data.get("imaging", {})
    p = data.get("pricing", {})
    shared = _shared_block(data)
    gp = p.get("gp_rate", 0)
    tax_rate = shared.get("tax_rate", 0.0825)

    imaging_amt = img.get("imaging_amt", 0)
    shipping = img.get("shipping", 0)
    labor_marked = (
        img.get("labor_days", 0) * img.get("labor_daily_rate", 1600) * (1 + gp)
    )
    labor_misc = img.get("labor_misc_amt", 0)
    travel = img.get("travel_adder_retail", 0)
    imaging_install = labor_marked + labor_misc + travel

    mid = shared.get("mid", {})
    inc_mat, inc_lab, _ = _mid_state(shared)
    mid_material = mid.get("material_amt", 0) if inc_mat else 0
    mid_labor = mid.get("labor_amt", 0) if inc_lab else 0

    material_total = imaging_amt + shipping + mid_material
    installation_total = imaging_install + mid_labor
    tax = material_total * tax_rate
    grand_total = material_total + installation_total + tax
    return {
        "imaging_amt": imaging_amt,
        "imaging_shipping": shipping,
        "imaging_install": imaging_install,
        "imaging_install_labor_marked": labor_marked,
        "imaging_install_misc": labor_misc,
        "imaging_install_travel": travel,
        "mid_material": mid_material,
        "mid_labor": mid_labor,
        "material_total": material_total,
        "installation_total": installation_total,
        "tax": tax,
        "grand_total": grand_total,
    }


def _combined_totals(data):
    mode = _get_mode(data)
    if mode == "imaging_only":
        return _imaging_totals(data)

    canopies = _canopies_list(data)
    shared = _shared_block(data)
    per_canopy = [_canopy_totals(c) for c in canopies]

    inc_mat, inc_lab, _ = _mid_state(shared)
    mid = shared.get("mid", {})
    price_sign_mat = mid.get("material_amt", 0) if inc_mat else 0
    price_sign_inst = mid.get("labor_amt", 0) if inc_lab else 0
    demo = shared.get("demo", {})
    demo_retail = demo.get("retail", 0) if demo.get("include") else 0

    if mid.get("material_amt") is None and inc_mat:
        price_sign_mat = data.get("pricing", {}).get("mid_material", 0)
    if mid.get("labor_amt") is None and inc_lab:
        price_sign_inst = data.get("pricing", {}).get("mid_labor", 0)

    canopy_mat_sum = sum(t["canopy_mat"] for t in per_canopy)
    canopy_inst_sum = sum(t["canopy_inst"] for t in per_canopy)
    branding_mat_sum = sum(t["branding_mat"] for t in per_canopy)
    material_total = canopy_mat_sum + branding_mat_sum + price_sign_mat
    installation_total = canopy_inst_sum + price_sign_inst + demo_retail
    tax = material_total * shared.get("tax_rate", 0.0825)
    grand_total = material_total + installation_total + tax

    return {
        "per_canopy": per_canopy,
        "price_sign_mat": price_sign_mat,
        "price_sign_inst": price_sign_inst,
        "demo_retail": demo_retail,
        "material_total": material_total,
        "installation_total": installation_total,
        "tax": tax,
        "grand_total": grand_total,
    }


def _section_totals(data):
    mode = _get_mode(data)
    if mode == "imaging_only":
        return _imaging_totals(data)
    combined = _combined_totals(data)
    if mode == "single_canopy" and combined.get("per_canopy"):
        t = combined["per_canopy"][0]
        return {
            "canopy_mat": t["canopy_mat"],
            "canopy_inst": t["canopy_inst"],
            "branding_mat": t["branding_mat"],
            "branding_inst": t["branding_inst"],
            "price_sign_mat": combined["price_sign_mat"],
            "price_sign_inst": combined["price_sign_inst"],
            "demo_retail": combined["demo_retail"],
            "material_total": combined["material_total"],
            "installation_total": combined["installation_total"],
            "tax": combined["tax"],
            "grand_total": combined["grand_total"],
        }
    return combined


# Inclusions / Exclusions --------------------------------------------------


def _derive_canopy_inclusions(canopy, *, label_prefix=None):
    light_type = canopy.get("light_type", "Standard")
    light_count = canopy.get("light_count", canopy.get("dispensers", 0) * 4)
    flood_count = canopy.get("flood_light_count", 0)
    inc = ["Canopy structure (ACM, steel, decking, miscellaneous)"]
    if light_type != "Customer-provided":
        inc.append(f"Canopy lighting — {light_count} fixtures ({light_type})")
    if flood_count > 0:
        inc.append(f"Flood lights — {flood_count} fixtures")
    inc += [
        "Engineering drawing (one per canopy)",
        "Canopy installation",
    ]
    customer_supplied = _customer_supplied_imaging(canopy)
    if canopy.get("branded"):
        if not customer_supplied:
            inc.append(f"Brand imaging — {canopy.get('brand_name', '')}")
            inc.append("Brand-imaging shipping & handling")
        else:
            inc.append(
                f"Brand imaging installation labor only "
                f"({canopy.get('brand_name', '')} material by customer)"
            )
        inc.append("Branding installation")
    else:
        inc.append("Forecourt painting (columns, island forms, bollards)")
    return inc


def _derive_canopy_exclusions(canopy):
    exc = []
    if canopy.get("light_type") == "Customer-provided":
        _lc = canopy.get("light_count", canopy.get("dispensers", 0) * 4)
        exc.append(f"Canopy lighting — {_lc} fixtures (to be provided by customer)")
    customer_supplied = _customer_supplied_imaging(canopy)
    if not canopy.get("branded"):
        exc.append("Brand imaging material & installation")
    elif customer_supplied:
        exc.append("Brand imaging material (to be provided by customer)")
    return exc


def _derive_shared_inclusions(data):
    shared = _shared_block(data)
    inc = []
    inc_mat, inc_lab, mid_brand = _mid_state(shared)
    if inc_mat and inc_lab:
        inc.append(f"Price sign — {mid_brand or ''}")
        inc.append("Price sign installation")
    elif inc_mat and not inc_lab:
        inc.append(f"Price sign material — {mid_brand or ''} (installation by others)")
    elif (not inc_mat) and inc_lab:
        inc.append(
            f"Price sign installation only (customer-supplied material"
            f"{f' — {mid_brand}' if mid_brand else ''})"
        )
    demo = shared.get("demo", {})
    if demo.get("include"):
        inc.append("Removal of existing canopy")
    return inc


def _derive_general_exclusions(data):
    exc = [
        "Permit fees",
        "Electrical work",
        "Canopy piers / foundations",
        "Bricking / decorative masonry",
        "Third-party inspections (provided at additional cost if required)",
    ]
    shared = _shared_block(data)
    inc_mat, inc_lab, mid_brand = _mid_state(shared)
    if not inc_mat and not inc_lab:
        exc.append("Price sign material & installation")
    elif inc_mat and not inc_lab:
        exc.append("Price sign installation (material only)")
    elif (not inc_mat) and inc_lab:
        exc.append("Price sign material (installation only; material by customer)")
    demo = shared.get("demo", {})
    if demo.get("include"):
        exc.append("Removal of canopy footers (footers stay in place)")
    canopies = _canopies_list(data)
    has_unbranded = any(not c.get("branded") for c in canopies)
    if (
        (inc_mat or inc_lab)
        and mid_brand
        and mid_brand != "Unbranded"
        and has_unbranded
    ):
        exc.append(
            "Canopy imaging (by others — scope on this proposal "
            "covers the price sign only on unbranded canopies)"
        )
    for extra in shared.get("items_not_included", []):
        if extra and extra not in exc:
            exc.append(extra)
    for c in canopies:
        for extra in c.get("pricing", {}).get("items_not_included", []):
            if extra and extra not in exc:
                exc.append(extra)
    return exc


def _derive_inclusions(data):
    inc = []
    canopies = _canopies_list(data)
    for c in canopies:
        inc += _derive_canopy_inclusions(c)
    inc += _derive_shared_inclusions(data)
    return inc


def _derive_exclusions(data):
    exc = list(_derive_general_exclusions(data))
    canopies = _canopies_list(data)
    for c in canopies:
        for e in _derive_canopy_exclusions(c):
            if e not in exc:
                exc.append(e)
    return exc


# Section primitives ------------------------------------------------------


def _add_section_band(doc, label):
    t = doc.add_table(rows=1, cols=1)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W])
    c = t.cell(0, 0)
    _shade_cell(c, _H(PALETTE["band_bg"]))
    _set_cell_margins(c, top=70, left=120, bottom=70, right=120)
    _clear_default_para(c)
    _add_para(c, label, size=10, bold=True, color=WHITE, space_after=0)


def _add_section_subtotal(doc, mat, inst, *, section_name="Section"):
    total = mat + inst
    t = doc.add_table(rows=1, cols=2)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W * 0.55, USABLE_W * 0.45])
    _shade_cell(t.cell(0, 0), _H(PALETTE["subtle_bg"]))
    _shade_cell(t.cell(0, 1), _H(PALETTE["subtle_bg"]))
    for ci in range(2):
        _set_cell_margins(t.cell(0, ci), top=60, bottom=60, left=160, right=160)
        _clear_default_para(t.cell(0, ci))
    _add_para(
        t.cell(0, 0),
        f"Material: {_money(mat)}     Installation: {_money(inst)}",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_para(
        t.cell(0, 1),
        f"{section_name} Total: {_money(total)}",
        size=11,
        bold=True,
        color=PALETTE["accent"],
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )


def _add_single_subtotal(doc, amount, *, label="Total"):
    t = doc.add_table(rows=1, cols=2)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W * 0.55, USABLE_W * 0.45])
    _shade_cell(t.cell(0, 0), _H(PALETTE["subtle_bg"]))
    _shade_cell(t.cell(0, 1), _H(PALETTE["subtle_bg"]))
    for ci in range(2):
        _set_cell_margins(t.cell(0, ci), top=60, bottom=60, left=160, right=160)
        _clear_default_para(t.cell(0, ci))
    _add_para(t.cell(0, 0), "", size=9, space_after=0)
    _add_para(
        t.cell(0, 1),
        f"{label}: {_money(amount)}",
        size=11,
        bold=True,
        color=PALETTE["accent"],
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )


def _add_section_subtotal_note(doc, note_text):
    t = doc.add_table(rows=1, cols=1)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W])
    cell = t.cell(0, 0)
    _shade_cell(cell, _H(PALETTE["subtle_bg"]))
    _set_cell_margins(cell, top=50, bottom=50, left=160, right=160)
    _clear_default_para(cell)
    _add_para(
        cell,
        note_text,
        size=9,
        bold=True,
        italic=True,
        color=PALETTE["text_emphasis"],
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )


# CANOPY + CANOPY BRANDING ------------------------------------------------


def _add_canopy_section(doc, canopy, totals, *, label_prefix=None):
    _add_section_band(doc, _canopy_band_label(canopy, label_prefix))
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _set_cell_margins(cell, top=100, left=160, bottom=80, right=160)
    _clear_default_para(cell)
    light_type = canopy.get("light_type", "Standard")
    light_count = canopy.get("light_count", canopy.get("dispensers", 0) * 4)
    flood_count = canopy.get("flood_light_count", 0)
    if light_type == "Customer-provided":
        lighting_line = "Canopy lighting — provided by customer (not included)"
    else:
        lighting_line = f"Canopy lighting ({light_count} fixtures, {light_type})"
    items_provided = [
        "ACM panels & fascia",
        f"Structural steel framing ({canopy.get('columns', 0)} columns)",
        "Decking system",
        lighting_line,
    ]
    if flood_count > 0:
        items_provided.append(f"Flood lights ({flood_count} fixtures)")
    items_provided.append(
        "Miscellaneous (paint, transport, project management, insurance)"
    )
    twocol = cell.add_table(rows=1, cols=2)
    _remove_table_borders(twocol)
    _set_fixed_layout(twocol, [3.45, 3.45])
    left_c, right_c = twocol.cell(0, 0), twocol.cell(0, 1)
    _set_cell_margins(left_c, top=0, bottom=0, left=0, right=120)
    _set_cell_margins(right_c, top=0, bottom=0, left=120, right=0)
    _clear_default_para(left_c)
    _clear_default_para(right_c)
    _add_para(
        left_c,
        "Items provided:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    for it in items_provided:
        _add_para(left_c, f"•  {it}", size=9, space_after=1)
    _add_para(
        right_c,
        "Build specifications:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    for spec in CANOPY_BUILD_SPECS:
        _add_para(right_c, f"•  {spec}", size=9, space_after=1)
    _add_para(cell, "", size=2, space_after=2)
    _add_para(
        cell,
        "Engineering drawing: One included per canopy.",
        size=9,
        italic=True,
        color=GREY,
        space_after=2,
    )
    _add_para(
        cell,
        "Third-party inspections: Not included; if required by the "
        "authority having jurisdiction, provided at additional cost.",
        size=9,
        italic=True,
        color=GREY,
        space_after=0,
    )
    if label_prefix and label_prefix.upper().endswith("CANOPY"):
        section_name = label_prefix.title()
    else:
        section_name = "Canopy"
    _add_section_subtotal(
        doc, totals["canopy_mat"], totals["canopy_inst"], section_name=section_name
    )


def _add_branding_section(doc, canopy, totals, *, label_prefix=None):
    branded = canopy.get("branded")
    customer_supplied = _customer_supplied_imaging(canopy)
    base = "CANOPY BRANDING" if not label_prefix else f"{label_prefix} BRANDING"
    if branded:
        if customer_supplied:
            band_label = (
                f"{base}  —  {canopy.get('brand_name', '')}  (Material by Customer)"
            )
        else:
            band_label = f"{base}  —  {canopy.get('brand_name', '')}"
    else:
        band_label = f"{base}  —  Unbranded"
    _add_section_band(doc, band_label)
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _set_cell_margins(cell, top=100, left=160, bottom=80, right=160)
    _clear_default_para(cell)
    if branded and not customer_supplied:
        _add_para(
            cell,
            IMAGING_DISCLAIMER,
            size=9,
            bold=True,
            italic=True,
            color=PALETTE["accent"],
            space_after=4,
        )
    _add_para(
        cell,
        "Items provided:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    if customer_supplied:
        items = ["Forecourt: paint of columns, island forms, and bollards"]
    else:
        items = [
            "Fascia / ACM imaging or paint",
            "Forecourt: paint of columns, island forms, and bollards",
        ]
    for it in items:
        _add_para(cell, f"•  {it}", size=9, space_after=1)
    if branded and not customer_supplied:
        for it in ("Brand valances", "Trash bins", "Pump toppers", "Flags"):
            _add_para(cell, f"•  {it}", size=9, space_after=1)
    _add_para(cell, "", size=2, space_after=2)
    if customer_supplied:
        concluding = (
            f"Imaging material for {canopy.get('brand_name', '')} to be provided by the "
            "customer — not quoted here. Installation labor for the "
            "customer-supplied imaging is included in the Canopy section."
        )
    elif branded:
        concluding = (
            "Branding retail prices include material (with shipping and "
            "handling). Installation is included in the Canopy section."
        )
    else:
        concluding = "Installation is included in the Canopy section."
    _add_para(cell, concluding, size=9, italic=True, color=GREY, space_after=0)
    if branded and not customer_supplied:
        section_name = (
            "Branding" if not label_prefix else label_prefix.title() + " Branding"
        )
        _add_section_subtotal(
            doc,
            totals["branding_mat"],
            totals["branding_inst"],
            section_name=section_name,
        )
    else:
        if customer_supplied:
            _add_section_subtotal_note(
                doc, "Material provided by customer — not included in this quote"
            )
        else:
            _add_section_subtotal_note(doc, "Included in canopy price")


# PRICE SIGN, EXISTING CANOPY REMOVAL, ADDITIONAL NOTES -------------------


def _add_price_sign_section(doc, data, totals):
    shared = _shared_block(data)
    inc_mat, inc_lab, brand = _mid_state(shared)
    brand = brand or ""
    mat_amt = totals.get("price_sign_mat", 0)
    inst_amt = totals.get("price_sign_inst", 0)

    if inc_mat and inc_lab:
        suffix = "(Material & Installation Included)"
    elif inc_mat and not inc_lab:
        suffix = "(Material Only — Installation by Others)"
    elif (not inc_mat) and inc_lab:
        suffix = "(Installation Only — Material by Customer)"
    else:
        suffix = "(Not Included)"

    if brand and (inc_mat or inc_lab):
        band_label = f"PRICE SIGN  —  {brand}  {suffix}"
    elif inc_mat or inc_lab:
        band_label = f"PRICE SIGN  —  {suffix}"
    else:
        band_label = "PRICE SIGN  —  Not Included"
    _add_section_band(doc, band_label)

    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _set_cell_margins(cell, top=100, left=160, bottom=80, right=160)
    _clear_default_para(cell)

    canopies = _canopies_list(data)
    has_unbranded = any(not c.get("branded") for c in canopies)
    if has_unbranded and brand and brand != "Unbranded" and (inc_mat or inc_lab):
        _add_para(
            cell,
            "Canopy imaging by others on the unbranded canopy; APEC scope "
            "on this proposal includes branded price sign supply and/or "
            "installation only for that canopy.",
            size=9,
            italic=True,
            color=PALETTE["accent"],
            space_after=2,
        )

    if inc_mat and inc_lab:
        _add_para(
            cell,
            "Includes price sign material and installation.",
            size=9,
            italic=True,
            color=GREY,
            space_after=0,
        )
        _add_section_subtotal(doc, mat_amt, inst_amt, section_name="Price Sign")
    elif inc_mat and not inc_lab:
        _add_para(
            cell,
            "Material only. Installation to be performed by others — not "
            "included in this proposal.",
            size=9,
            italic=True,
            color=GREY,
            space_after=0,
        )
        _add_single_subtotal(doc, mat_amt, label="Material")
    elif (not inc_mat) and inc_lab:
        _add_para(
            cell,
            "Installation only. Price sign material to be supplied by "
            "customer; the amount below covers installation labor.",
            size=9,
            italic=True,
            color=GREY,
            space_after=0,
        )
        _add_single_subtotal(doc, inst_amt, label="Installation")
    else:
        _add_para(
            cell,
            "No price sign material or installation quoted here.",
            size=9,
            italic=True,
            color=GREY,
            space_after=0,
        )
        _add_section_subtotal_note(doc, "Not included in this quote")


def _add_demolition_section(doc, data, totals):
    shared = _shared_block(data)
    demo = shared.get("demo", {})
    if not demo.get("include"):
        return
    _add_section_band(doc, "EXISTING CANOPY REMOVAL")
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _set_cell_margins(cell, top=100, left=160, bottom=80, right=160)
    _clear_default_para(cell)
    _add_para(
        cell,
        "Scope of work:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    canopies = _canopies_list(data)
    plural = "canopies" if len(canopies) > 1 else "canopy"
    bullets = [
        f"Remove existing {plural}",
        "Remove and discard all old canopy material",
        "Columns to be cut at the bottom down to concrete level",
        "Footers removal is not covered under this quote",
    ]
    for b in bullets:
        _add_para(cell, f"•  {b}", size=9, space_after=1)
    _add_para(cell, "", size=2, space_after=2)
    _add_single_subtotal(
        doc, totals.get("demo_retail", 0), label="Existing Canopy Removal"
    )


def _add_notes_box(doc):
    _add_section_band(doc, "ADDITIONAL NOTES")
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _shade_cell(cell, "FFFDF5")
    _set_cell_margins(cell, top=120, left=160, bottom=600, right=160)
    _clear_default_para(cell)
    for _ in range(3):
        _add_para(cell, "", size=10, space_after=0)


# INCLUDED / NOT INCLUDED -------------------------------------------------


def _add_subhead(cell, text, color):
    _add_para(cell, text, size=9, bold=True, color=color, space_after=2, space_before=4)


def _add_incl_excl(doc, data):
    incexc = doc.add_table(rows=1, cols=2)
    _remove_table_borders(incexc)
    _set_fixed_layout(incexc, [USABLE_W / 2, USABLE_W / 2])
    inc_cell, exc_cell = incexc.cell(0, 0), incexc.cell(0, 1)
    _shade_cell(inc_cell, "EAF7EE")
    _shade_cell(exc_cell, "FCEEEE")
    _set_cell_margins(inc_cell, top=120, left=140, bottom=120, right=140)
    _set_cell_margins(exc_cell, top=120, left=140, bottom=120, right=140)
    _clear_default_para(inc_cell)
    _clear_default_para(exc_cell)
    _add_para(inc_cell, "✓  INCLUDED", size=10, bold=True, color=GREEN, space_after=4)
    _add_para(exc_cell, "✗  NOT INCLUDED", size=10, bold=True, color=RED, space_after=4)

    canopies = _canopies_list(data)
    multi = len(canopies) > 1

    for c in canopies:
        items = _derive_canopy_inclusions(c)
        if multi:
            label = f"{c.get('fuel_type', '').upper()} CANOPY".strip()
            _add_subhead(inc_cell, label, PALETTE["text_emphasis"])
        for it in items:
            _add_para(
                inc_cell,
                f"•  {it}",
                size=9,
                color=PALETTE["text_emphasis"],
                space_after=1,
            )
    shared_inc = _derive_shared_inclusions(data)
    if shared_inc:
        if multi:
            _add_subhead(inc_cell, "SHARED", PALETTE["text_emphasis"])
        for it in shared_inc:
            _add_para(
                inc_cell,
                f"•  {it}",
                size=9,
                color=PALETTE["text_emphasis"],
                space_after=1,
            )

    any_canopy_specific = False
    for c in canopies:
        per_exc = _derive_canopy_exclusions(c)
        if per_exc:
            any_canopy_specific = True
            if multi:
                label = f"{c.get('fuel_type', '').upper()} CANOPY".strip()
                _add_subhead(exc_cell, label, PALETTE["text_emphasis"])
            for it in per_exc:
                _add_para(
                    exc_cell,
                    f"•  {it}",
                    size=9,
                    color=PALETTE["text_emphasis"],
                    space_after=1,
                )
    general = _derive_general_exclusions(data)
    if general:
        if multi or any_canopy_specific:
            _add_subhead(exc_cell, "GENERAL", PALETTE["text_emphasis"])
        for it in general:
            _add_para(
                exc_cell,
                f"•  {it}",
                size=9,
                color=PALETTE["text_emphasis"],
                space_after=1,
            )


# IMAGING-only section ----------------------------------------------------


def _add_imaging_section(doc, data, totals):
    img = data.get("imaging", {})
    brand = img.get("brand") or "Unbranded"
    customer_supplied = bool(img.get("customer_supplied_imaging"))
    if customer_supplied:
        band_label = f"IMAGING  —  {brand}  (Material by Customer)"
    else:
        band_label = f"IMAGING  —  {brand}"
    _add_section_band(doc, band_label)
    body = doc.add_table(rows=1, cols=1)
    _remove_table_borders(body)
    _set_fixed_layout(body, [USABLE_W])
    cell = body.cell(0, 0)
    _set_cell_margins(cell, top=100, left=160, bottom=80, right=160)
    _clear_default_para(cell)
    if not customer_supplied:
        _add_para(
            cell,
            IMAGING_DISCLAIMER,
            size=9,
            bold=True,
            italic=True,
            color=PALETTE["accent"],
            space_after=4,
        )
    _add_para(
        cell,
        "Scope of work:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    if customer_supplied:
        bullets = [
            f"Installation of customer-supplied {brand} imaging material",
            "Forecourt: paint of columns, island forms, and bollards",
            "Site preparation and clean-up",
        ]
    else:
        bullets = [
            f"Brand imaging material — {brand}",
            "Brand-imaging shipping & handling",
            "Fascia / ACM imaging or paint",
            "Forecourt: paint of columns, island forms, and bollards",
            "Brand valances, trash bins, pump toppers, flags",
            "Imaging installation labor",
        ]
    for b in bullets:
        _add_para(cell, f"•  {b}", size=9, space_after=1)
    _add_para(cell, "", size=2, space_after=2)
    days = img.get("labor_days", 0)
    misc = img.get("labor_misc_amt", 0)
    _add_para(
        cell,
        f"Installation labor: {days} days  +  miscellaneous {_money(misc)}.",
        size=9,
        italic=True,
        color=GREY,
        space_after=0,
    )
    if customer_supplied:
        mat_amt = 0
    else:
        mat_amt = totals.get("imaging_amt", 0) + totals.get("imaging_shipping", 0)
    inst_amt = totals.get("imaging_install", 0)
    _add_section_subtotal(doc, mat_amt, inst_amt, section_name="Imaging")


# GRAND TOTAL, TERMS, SIGNATURE ------------------------------------------


def _add_grand_total(doc, totals, tax_rate):
    t = doc.add_table(rows=4, cols=2)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W - 2.0, 2.0])
    rows = [
        ("Material total", _money(totals["material_total"]), False),
        ("Installation", _money(totals["installation_total"]), False),
        (
            f"Sales tax ({tax_rate * 100:.2f}% on material)",
            _money(totals["tax"]),
            False,
        ),
        ("GRAND TOTAL", _money(totals["grand_total"]), True),
    ]
    for ri, (label, amount, is_grand) in enumerate(rows):
        l, r = t.cell(ri, 0), t.cell(ri, 1)
        _set_cell_margins(l, top=60, bottom=60, left=160)
        _set_cell_margins(r, top=60, bottom=60, right=160)
        _clear_default_para(l)
        _clear_default_para(r)
        if is_grand:
            _shade_cell(l, _H(PALETTE["grand_total_bg"]))
            _shade_cell(r, _H(PALETTE["grand_total_bg"]))
            _add_para(l, label, size=12, bold=True, color=WHITE, space_after=0)
            _add_para(
                r,
                amount,
                size=14,
                bold=True,
                color=WHITE,
                space_after=0,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )
        else:
            _shade_cell(l, _H(PALETTE["subtle_bg"]))
            _shade_cell(r, _H(PALETTE["subtle_bg"]))
            _add_para(
                l,
                label,
                size=10,
                bold=False,
                color=PALETTE["text_emphasis"],
                space_after=0,
            )
            _add_para(
                r,
                amount,
                size=11,
                bold=True,
                color=PALETTE["text_emphasis"],
                space_after=0,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )


def _add_terms(doc):
    t = doc.add_table(rows=1, cols=1)
    _remove_table_borders(t)
    _set_fixed_layout(t, [USABLE_W])
    cell = t.cell(0, 0)
    _shade_cell(cell, _H(PALETTE["subtle_bg"]))
    _set_cell_margins(cell, top=100, left=160, bottom=100, right=160)
    _clear_default_para(cell)
    _add_para(
        cell,
        "Disclaimers:",
        size=9,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=3,
    )
    disclosures = (
        "Pricing is contingent on jobsite accessibility and favorable "
        "ground conditions. "
        "All material cost has been included on an estimated basis using "
        "certain assumptions. This may be changed based on the prevailing "
        "prices and actual quantities at the time of the order and hence "
        "will be charged to the customer separately. "
        "Material warranty is based on manufacturers’ warranties while "
        "installation warranty is 90 days from the date of installation. "
        "Any requirement such as soil testing, electrical, survey drawings, "
        "third party or Govt. special inspections costs are not included "
        "and if needed for permitting, will be done at owner’s expense. "
        "Carrying charges of 1.5% per month will be charged on all past "
        "due payments."
    )
    _add_para(cell, disclosures, size=7, color=GREY, space_after=0)


def _add_signature_block(doc, company):
    sig = doc.add_table(rows=2, cols=2)
    _remove_table_borders(sig)
    _set_fixed_layout(sig, [USABLE_W / 2, USABLE_W / 2])
    for ci in range(2):
        line_cell = sig.cell(0, ci)
        _set_cell_margins(line_cell, top=280, bottom=20)
        _clear_default_para(line_cell)
        _add_cell_bottom_border(line_cell)
        _add_para(line_cell, "", size=8, space_after=0)
    _clear_default_para(sig.cell(1, 0))
    _clear_default_para(sig.cell(1, 1))
    _add_para(
        sig.cell(1, 0),
        "Customer Signature, Name & Date",
        size=8,
        color=GREY,
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_para(
        sig.cell(1, 1),
        company["signature_label"],
        size=8,
        color=GREY,
        space_after=0,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )


# Project Specifications --------------------------------------------------


def _spec_rows_for_canopy(canopy):
    col_arr = (
        "Double column"
        if (canopy["type"] == "Dive-in" and canopy.get("double_col"))
        else "Single column"
    )
    canopy_label = (
        f"{canopy['type']} ({col_arr})"
        if canopy["type"] == "Dive-in"
        else canopy["type"]
    )
    branded_label = (
        f"Yes — {canopy.get('brand_name', '')}" if canopy.get("branded") else "No"
    )
    lo = canopy.get("left_overhang", canopy.get("overhang", 0))
    ro = canopy.get("right_overhang", canopy.get("overhang", 0))
    sp = canopy.get("spacing", canopy.get("distance", 0))
    spacing_row_label = f"{int(lo)}′  /  {int(sp)}′  /  {int(ro)}′"
    rows = [
        ("Fuel Type", canopy.get("fuel_type", "—")),
        ("Canopy Type", canopy_label),
        ("Branding", branded_label),
        ("Dispensers / Columns", f"{canopy['dispensers']}  /  {canopy['columns']}"),
        ("Dimensions (W × D)", f"{int(canopy['width'])}′  ×  {int(canopy['depth'])}′"),
        ("Square Footage", f"{int(canopy['width'] * canopy['depth']):,} sq ft"),
        ("LO / Spacing / RO", spacing_row_label),
        ("Lights", f"{canopy.get('light_count', canopy['dispensers'] * 4)} fixtures"),
        ("Light Type", canopy.get("light_type", "Standard")),
    ]
    flood_count = canopy.get("flood_light_count", 0)
    if flood_count > 0:
        rows.append(("Flood Lights", f"{flood_count} fixtures"))
    return rows


def _spec_rows(data):
    canopies = _canopies_list(data)
    canopy = canopies[0]
    rows = _spec_rows_for_canopy(canopy)
    inc_mat, inc_lab, mid_brand = _mid_state(data)
    if inc_mat and inc_lab:
        price_sign_label = f"Yes — {mid_brand or ''}"
    elif inc_mat and not inc_lab:
        price_sign_label = f"Material only — {mid_brand or ''}"
    elif (not inc_mat) and inc_lab:
        price_sign_label = (
            f"Installation only — {mid_brand or ''}"
            if mid_brand
            else "Installation only"
        )
    else:
        price_sign_label = "No"
    inserted = []
    for k, v in rows:
        inserted.append((k, v))
        if k == "Branding":
            inserted.append(("Price Sign", price_sign_label))
    return inserted


def _spec_block_caption(canopy_or_data):
    if "canopy" in canopy_or_data:
        canopy = canopy_or_data["canopy"]
    elif "fuel_type" in canopy_or_data or "type" in canopy_or_data:
        canopy = canopy_or_data
    else:
        canopies = _canopies_list(canopy_or_data)
        canopy = canopies[0]
    fuel = canopy.get("fuel_type", "")
    if canopy.get("branded"):
        brand_part = f"Branded — {canopy.get('brand_name', '')}"
        if _customer_supplied_imaging(canopy):
            brand_part += " (Material by Customer)"
    else:
        brand_part = "Unbranded"
    if fuel:
        return f"{fuel}  •  {brand_part}"
    return brand_part


# build_proposal ----------------------------------------------------------


def _build_header(doc, company):
    header = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header)
    _set_fixed_layout(header, [2.1, USABLE_W - 2.1])
    logo_cell = header.cell(0, 0)
    info_cell = header.cell(0, 1)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_default_para(logo_cell)
    logo_path = Path(__file__).parent / company["logo"]
    if logo_path.exists():
        p = logo_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(logo_path), width=Inches(1.8))
    _clear_default_para(info_cell)
    _add_para(
        info_cell,
        company["name"],
        size=14,
        bold=True,
        color=PALETTE["accent"],
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=1,
    )
    _add_para(
        info_cell,
        company["tagline"],
        size=9,
        italic=True,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=4,
    )
    _add_para(
        info_cell,
        company["address"],
        size=9,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=1,
    )
    _add_para(
        info_cell,
        f"{company['phone']}  •  {company['email']}  •  {company['website']}",
        size=9,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=0,
    )


def _build_title(doc, data):
    title = doc.add_table(rows=1, cols=3)
    _remove_table_borders(title)
    _set_fixed_layout(title, [2.7, 2.4, USABLE_W - 5.1])
    for ci in range(3):
        _shade_cell(title.cell(0, ci), _H(PALETTE["band_bg"]))
        _set_cell_margins(title.cell(0, ci), top=70, bottom=70)
    c0, c1, c2 = title.cell(0, 0), title.cell(0, 1), title.cell(0, 2)
    _clear_default_para(c0)
    _clear_default_para(c1)
    _clear_default_para(c2)
    _add_para(c0, "PROPOSAL", size=15, bold=True, color=WHITE, space_after=0)
    p1 = c1.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_run(p1.add_run("QUOTE #  "), size=9, color=WHITE)
    _format_run(p1.add_run(data["quote_number"]), size=11, bold=True, color=WHITE)
    p2 = c2.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_run(p2.add_run("DATE  "), size=9, color=WHITE)
    _format_run(p2.add_run(data["quote_date"]), size=11, bold=True, color=WHITE)
    _add_para(doc, "", size=2, space_after=0)


def _render_spec_table(parent_cell, rows, col_widths):
    spec_tbl = parent_cell.add_table(rows=len(rows), cols=2)
    _remove_table_borders(spec_tbl)
    _set_fixed_layout(spec_tbl, col_widths)
    for ri, (k, v) in enumerate(rows):
        kcell, vcell = spec_tbl.cell(ri, 0), spec_tbl.cell(ri, 1)
        if ri % 2 == 0:
            _shade_cell(kcell, _H(PALETTE["subtle_bg_alt"]))
            _shade_cell(vcell, _H(PALETTE["subtle_bg_alt"]))
        _set_cell_margins(kcell, top=30, bottom=30, left=80, right=20)
        _set_cell_margins(vcell, top=30, bottom=30, left=20, right=80)
        _clear_default_para(kcell)
        _clear_default_para(vcell)
        _add_para(kcell, k, size=9, color=GREY, space_after=0)
        _add_para(
            vcell,
            v,
            size=9,
            bold=True,
            color=PALETTE["text_emphasis"],
            space_after=0,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )


def _build_customer_specs(doc, data):
    cust_specs = doc.add_table(rows=1, cols=2)
    _remove_table_borders(cust_specs)
    _set_fixed_layout(cust_specs, [3.3, USABLE_W - 3.3])
    left = cust_specs.cell(0, 0)
    _clear_default_para(left)
    _shade_cell(left, _H(PALETTE["customer_bg"]))
    _set_cell_margins(left, top=120, left=140, bottom=120, right=140)
    cust = data["customer"]
    _add_para(
        left,
        "PREPARED FOR",
        size=8,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    _add_para(
        left,
        cust.get("company", ""),
        size=11,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=1,
    )
    if cust.get("name"):
        _add_para(left, f"Attn: {cust['name']}", size=9, color=GREY, space_after=1)
    contact_bits = [b for b in (cust.get("phone"), cust.get("email")) if b]
    if contact_bits:
        _add_para(left, "  •  ".join(contact_bits), size=9, color=GREY, space_after=8)
    else:
        _add_para(left, "", space_after=8)
    _add_para(
        left,
        "PROJECT SITE",
        size=8,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    _add_para(left, cust.get("street", ""), size=10, space_after=1)
    citystatezip = ", ".join(
        b
        for b in (
            cust.get("city"),
            f"{cust.get('state', '')} {cust.get('zip', '')}".strip(),
        )
        if b and b.strip()
    )
    _add_para(left, citystatezip, size=10, space_after=6)
    _add_para(
        left, "TERMS", size=8, bold=True, color=PALETTE["text_emphasis"], space_after=2
    )
    _add_para(left, "Quote valid for 30 days", size=10, space_after=1)
    _add_para(
        left, "Payment: 50% deposit / 50% at start of work", size=10, space_after=6
    )
    sp = data.get("sales_person", {})
    if sp.get("name"):
        _add_para(
            left,
            "SALES REPRESENTATIVE",
            size=8,
            bold=True,
            color=PALETTE["text_emphasis"],
            space_after=2,
        )
        _add_para(left, sp.get("name", ""), size=10, bold=True, space_after=1)
        sp_contact = [b for b in (sp.get("phone"), sp.get("email")) if b]
        if sp_contact:
            _add_para(left, "  •  ".join(sp_contact), size=9, color=GREY, space_after=0)

    right = cust_specs.cell(0, 1)
    _clear_default_para(right)
    _set_cell_margins(right, top=120, left=140, bottom=120, right=140)
    _add_para(
        right,
        "PROJECT SPECIFICATIONS",
        size=8,
        bold=True,
        color=PALETTE["text_emphasis"],
        space_after=2,
    )
    mode = _get_mode(data)
    if mode == "imaging_only":
        img = data.get("imaging", {})
        _add_para(
            right,
            f"Imaging Only  •  {img.get('brand', '—')}",
            size=10,
            bold=True,
            color=PALETTE["accent"],
            space_after=4,
        )
        rows = [
            ("Quote Type", "Imaging Only"),
            ("Brand", img.get("brand", "—")),
            ("Dispensers", str(img.get("dispensers", "—"))),
            (
                "Imaging Supplier",
                "Customer"
                if img.get("customer_supplied_imaging")
                else f"{data.get('company_key', 'APEC')} provides",
            ),
            ("Labor Days", str(img.get("labor_days", 0))),
            ("Misc Labor", _money(img.get("labor_misc_amt", 0))),
        ]
        inc_mat, inc_lab, mid_brand = _mid_state(data)
        if inc_mat or inc_lab:
            if inc_mat and inc_lab:
                ps_label = f"Yes — {mid_brand or ''}"
            elif inc_mat:
                ps_label = f"Material only — {mid_brand or ''}"
            else:
                ps_label = f"Installation only — {mid_brand or ''}"
            rows.append(("Price Sign", ps_label))
        _render_spec_table(right, rows, [1.7, 1.7])
    elif mode == "double_canopy":
        canopies = _canopies_list(data)
        labels = [f"{c.get('fuel_type', '').upper()} CANOPY".strip() for c in canopies]
        _add_para(
            right,
            "  |  ".join(_spec_block_caption(c) for c in canopies),
            size=10,
            bold=True,
            color=PALETTE["accent"],
            space_after=4,
        )
        rows_a = _spec_rows_for_canopy(canopies[0])
        rows_b = _spec_rows_for_canopy(canopies[1])
        n = max(len(rows_a), len(rows_b))
        spec_tbl = right.add_table(rows=n + 1, cols=4)
        _remove_table_borders(spec_tbl)
        col_w = (USABLE_W - 3.3 - 0.4) / 4
        _set_fixed_layout(spec_tbl, [col_w] * 4)
        for ci, lab in enumerate(labels):
            hcell = spec_tbl.cell(0, ci * 2)
            hcell2 = spec_tbl.cell(0, ci * 2 + 1)
            for hc in (hcell, hcell2):
                _shade_cell(hc, _H(PALETTE["subtle_bg_alt"]))
                _set_cell_margins(hc, top=20, bottom=20, left=60, right=20)
                _clear_default_para(hc)
            _add_para(
                hcell,
                lab,
                size=8,
                bold=True,
                color=PALETTE["text_emphasis"],
                space_after=0,
            )
            _add_para(hcell2, "", size=8, space_after=0)
        for ri in range(n):
            for side, rows in enumerate((rows_a, rows_b)):
                if ri < len(rows):
                    k, v = rows[ri]
                else:
                    k, v = "", ""
                kcell = spec_tbl.cell(ri + 1, side * 2)
                vcell = spec_tbl.cell(ri + 1, side * 2 + 1)
                if (ri + 1) % 2 == 0:
                    _shade_cell(kcell, _H(PALETTE["subtle_bg_alt"]))
                    _shade_cell(vcell, _H(PALETTE["subtle_bg_alt"]))
                _set_cell_margins(kcell, top=20, bottom=20, left=60, right=10)
                _set_cell_margins(vcell, top=20, bottom=20, left=10, right=40)
                _clear_default_para(kcell)
                _clear_default_para(vcell)
                _add_para(kcell, k, size=8, color=GREY, space_after=0)
                _add_para(
                    vcell,
                    v,
                    size=8,
                    bold=True,
                    color=PALETTE["text_emphasis"],
                    space_after=0,
                    align=WD_ALIGN_PARAGRAPH.RIGHT,
                )
        shared = _shared_block(data)
        inc_mat, inc_lab, mid_brand = _mid_state(shared)
        if inc_mat and inc_lab:
            ps = f"Yes — {mid_brand or ''}"
        elif inc_mat and not inc_lab:
            ps = f"Material only — {mid_brand or ''}"
        elif (not inc_mat) and inc_lab:
            ps = (
                f"Installation only — {mid_brand or ''}"
                if mid_brand
                else "Installation only"
            )
        else:
            ps = "No"
        demo_lbl = "Yes" if shared.get("demo", {}).get("include") else "No"
        _add_para(right, "", size=2, space_after=0)
        _add_para(
            right,
            "SHARED",
            size=8,
            bold=True,
            color=PALETTE["text_emphasis"],
            space_after=2,
        )
        _render_spec_table(
            right,
            [("Price Sign", ps), ("Existing Canopy Removal", demo_lbl)],
            [(USABLE_W - 3.3 - 0.4) / 2] * 2,
        )
    else:
        _add_para(
            right,
            _spec_block_caption(data),
            size=10,
            bold=True,
            color=PALETTE["accent"],
            space_after=4,
        )
        spec_rows = _spec_rows(data)
        _render_spec_table(right, spec_rows, [1.7, 1.7])
    _add_para(doc, "", size=2, space_after=0)


def build_proposal(data, output_path):
    global PALETTE
    output_path = Path(output_path)
    company = COMPANIES[data["company_key"]]
    PALETTE = company["palette"]
    totals = _combined_totals(data)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _build_header(doc, company)
    _build_title(doc, data)
    _build_customer_specs(doc, data)

    mode = _get_mode(data)
    if mode == "imaging_only":
        _add_imaging_section(doc, data, totals)
        _add_para(doc, "", size=2, space_after=0)
        _add_price_sign_section(doc, data, totals)
        _add_para(doc, "", size=2, space_after=0)
    elif mode == "double_canopy":
        canopies = _canopies_list(data)
        per_canopy = totals["per_canopy"]
        labels = [f"{c.get('fuel_type', '').upper()} CANOPY".strip() for c in canopies]
        for canopy, ct, lbl in zip(canopies, per_canopy, labels):
            _add_canopy_section(doc, canopy, ct, label_prefix=lbl)
            _add_para(doc, "", size=2, space_after=0)
            _add_branding_section(doc, canopy, ct, label_prefix=lbl)
            _add_para(doc, "", size=2, space_after=0)
        _add_price_sign_section(doc, data, totals)
        _add_para(doc, "", size=2, space_after=0)
        _add_demolition_section(doc, data, totals)
        if _shared_block(data).get("demo", {}).get("include"):
            _add_para(doc, "", size=2, space_after=0)
    else:
        canopies = _canopies_list(data)
        per_canopy = totals["per_canopy"]
        _add_canopy_section(doc, canopies[0], per_canopy[0], label_prefix=None)
        _add_para(doc, "", size=2, space_after=0)
        _add_branding_section(doc, canopies[0], per_canopy[0], label_prefix=None)
        _add_para(doc, "", size=2, space_after=0)
        _add_price_sign_section(doc, data, totals)
        _add_para(doc, "", size=2, space_after=0)
        _add_demolition_section(doc, data, totals)
        if _shared_block(data).get("demo", {}).get("include"):
            _add_para(doc, "", size=2, space_after=0)

    _add_notes_box(doc)
    _add_para(doc, "", size=2, space_after=0)
    if mode != "imaging_only":
        _add_incl_excl(doc, data)
        _add_para(doc, "", size=2, space_after=0)
    _add_grand_total(doc, totals, _shared_block(data).get("tax_rate", 0.0825))
    _add_para(doc, "", size=2, space_after=0)
    _add_terms(doc)
    _add_para(doc, "", size=2, space_after=0)
    _add_signature_block(doc, company)

    doc.save(str(output_path))
    return output_path


def sample_data():
    return {
        "company_key": "APEC",
        "quote_number": f"Q-{_dt.date.today().strftime('%Y%m%d')}-001",
        "quote_date": _dt.date.today().strftime("%B %d, %Y"),
        "customer": {
            "company": "Burleson Fuel Stop",
            "name": "Test Customer",
            "phone": "(817) 555-0100",
            "email": "owner@burlesonfuelstop.test",
            "street": "898 NE Alsbury Blvd",
            "city": "Burleson",
            "state": "TX",
            "zip": "76028",
        },
        "sales_person": {
            "name": "Walid Husain",
            "phone": "(555) 555-0123",
            "email": "walid@apecimaging.com",
        },
        "canopy": {
            "fuel_type": "Gas",
            "type": "Dive-in",
            "branded": False,
            "brand_name": "",
            "customer_supplied_imaging": False,
            "double_col": False,
            "dispensers": 4,
            "columns": 4,
            "width": 94,
            "depth": 24,
            "left_overhang": 10,
            "right_overhang": 10,
            "spacing": 28,
            "labor_days": 7,
            "light_type": "Standard",
            "light_count": 16,
            "flood_light_count": 0,
        },
        "mid": {"include_material": False, "include_labor": False, "brand": None},
        "demo": {"include": False, "retail": 0, "cost": 0},
        "pricing": {
            "gp_rate": 0.45,
            "tax_rate": 0.0825,
            "acm": 11_700,
            "steel": 16_000,
            "decking": 20_000,
            "lights": 2_400,
            "flood_lights": 0,
            "misc": 2_500,
            "steel_secondary": 0,
            "shipping": 0,
            "brand_imaging": 0,
            "mid_material": 0,
            "base_labor": 8_400,
            "branded_labor_add": 0,
            "two_disp_labor_add": 0,
            "mid_labor": 0,
            "travel_adder_retail": 0,
            "items_not_included": [],
        },
    }


if __name__ == "__main__":
    out = build_proposal(
        sample_data(), Path(__file__).parent / "Proposal_v13_Sample.docx"
    )
    print(f"Wrote {out}")
